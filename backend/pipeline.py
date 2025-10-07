# pipeline.py
"""
Pipeline that:
- extracts from files,
- maps lines to canonical params,
- extracts value+unit,
- stores raw extractions,
- persists ALL variants to MasterSpec (one row per variant),
- returns parsed_by_source and merged_master (for UI consumption).
"""

import os
import re
import json
import logging
import shutil
import time
from pathlib import Path
from typing import Dict, Any, List
import pdfplumber
from docx import Document
import io
import datetime
from PIL import Image
import pytesseract
from sentence_transformers import SentenceTransformer, util
import tempfile

from db import SessionLocal, engine, Base
from models import MasterSpec, RawExtraction
import pandas as pd
from s3_utils import download_prefix, upload_folder, download_file_stream, list_objects, upload_file

# ensure DB tables
Base.metadata.create_all(bind=engine)

ROOT = Path(__file__).parent
UPLOAD_DIR = ROOT / "data" / "uploads"
LANDING_DIR = ROOT / "data" / "landing"
OUTPUT_DIR = ROOT / "data" / "outputs"
LANDING_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Load defect rules once
DEFECTS_CSV_PATH = ROOT / "data" / "defects.csv"
DEFECT_RULES = pd.read_csv(DEFECTS_CSV_PATH).to_dict(orient="records")

logger = logging.getLogger("pipeline")
logging.basicConfig(level=logging.INFO)

# canonical params
CANONICAL = {
    "cap_diameter": ["cap diameter", "cap_dia", "cap diameter", "cap dia"],
    "tear_size_limit": ["tear size limit", "tear limit", "tear_size"],
    "surface_finish_tolerance": ["surface finish tolerance", "surface finish tol"],
    "hole_diameter": ["hole diameter", "hole dia"],
    "length_tolerance": ["length tolerance"],
    "width_tolerance": ["width tolerance"],
    "thickness_tolerance": ["thickness tolerance"],
    "material_type": ["material type", "material"],
    "max_pressure": ["max pressure", "operating pressure"],
    "max_temperature": ["max temperature", "max temp"],
    "min_temperature": ["min temperature", "min temp"],
}

logger.info("Loading embedding model...")
try:
    EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
    PARAM_EMBEDS = {k: EMBED_MODEL.encode(v, convert_to_tensor=True) for k, v in CANONICAL.items()}
    logger.info(f"Model loaded successfully. Canonical params: {list(CANONICAL.keys())}")
except Exception as e:
    logger.error(f"Failed to load embedding model: {e}")
    EMBED_MODEL = None
    PARAM_EMBEDS = {}


def extract_from_pdf(path: Path) -> str:
    text = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text).strip()


def extract_from_docx(path: Path) -> str:
    doc = Document(str(path))
    pieces = []
    for p in doc.paragraphs:
        if p.text.strip():
            pieces.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            row_text = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_text:
                pieces.append("\t".join(row_text))
    return "\n".join(pieces).strip()


def extract_from_image(path: Path) -> str:
    img = Image.open(str(path))
    txt = pytesseract.image_to_string(img)
    return txt.strip()


def extract_text_for_file(file_path: Path):
    suf = file_path.suffix.lower()
    if suf in (".pdf",):
        return extract_from_pdf(file_path)
    elif suf in (".docx",):
        return extract_from_docx(file_path)
    elif suf in (".png", ".jpg", ".jpeg", ".tiff"):
        return extract_from_image(file_path)
    else:
        return file_path.read_text(encoding="utf-8", errors="ignore")

def extract_text_for_s3_stream(file_stream, filename: str):
    """Extract text from S3 file stream based on file extension."""
    suf = Path(filename).suffix.lower()
    
    # For PDF and DOCX files, we need to save to temporary file since the libraries require file paths
    if suf in (".pdf", ".docx"):
        temp_file_path = None
        try:
            # Create temporary file
            suffix = ".pdf" if suf == ".pdf" else ".docx"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
                # Read the file stream content
                file_content = file_stream.read()
                temp_file.write(file_content)
                temp_file.flush()
                temp_file_path = temp_file.name
            
            # Close the temp file handle before processing
            # This is important on Windows to avoid file locking issues
            
            # Extract text from the file
            if suf == ".pdf":
                text = extract_from_pdf(Path(temp_file_path))
            else:  # .docx
                text = extract_from_docx(Path(temp_file_path))
            
            # Clean up temp file with retry logic for Windows
            _safe_delete_file(temp_file_path)
            return text
            
        except Exception as e:
            logger.error(f"Error processing {suf.upper()} file {filename}: {e}")
            if temp_file_path:
                _safe_delete_file(temp_file_path)
            return ""
            
    elif suf in (".png", ".jpg", ".jpeg", ".tiff"):
        try:
            image = Image.open(io.BytesIO(file_stream.read()))
            txt = pytesseract.image_to_string(image)
            return txt.strip()
        except Exception as e:
            logger.error(f"Error processing image file {filename}: {e}")
            return ""
    else:
        # For text files, read directly from stream
        try:
            return file_stream.read().decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error processing text file {filename}: {e}")
            return ""


def _safe_delete_file(file_path: str, max_retries: int = 3, delay: float = 0.1):
    """Safely delete a file with retry logic for Windows file locking issues."""
    if not file_path or not os.path.exists(file_path):
        return
    
    for attempt in range(max_retries):
        try:
            os.unlink(file_path)
            logger.debug(f"Successfully deleted temp file: {file_path}")
            return
        except (OSError, PermissionError) as e:
            if attempt < max_retries - 1:
                logger.warning(f"Failed to delete temp file {file_path} (attempt {attempt + 1}): {e}. Retrying...")
                time.sleep(delay)
            else:
                logger.error(f"Failed to delete temp file {file_path} after {max_retries} attempts: {e}")
                # On Windows, we might need to schedule the file for deletion on next reboot
                try:
                    import ctypes
                    if os.name == 'nt':  # Windows
                        ctypes.windll.kernel32.MoveFileExW(
                            file_path, None, 
                            ctypes.c_int(4)  # MOVEFILE_DELAY_UNTIL_REBOOT
                        )
                        logger.info(f"Scheduled temp file for deletion on reboot: {file_path}")
                except Exception:
                    pass


def map_line_to_param(line: str):
    if EMBED_MODEL is None or not PARAM_EMBEDS:
        logger.error("Embedding model not loaded. Cannot map line to param.")
        return None, 0.0
    
    try:
        emb = EMBED_MODEL.encode(line, convert_to_tensor=True)
        best_param = None
        best_score = -1.0
        for param, embeds in PARAM_EMBEDS.items():
            score = util.cos_sim(emb, embeds).max().item()
            if score > best_score:
                best_param = param
                best_score = score
        return best_param, best_score
    except Exception as e:
        logger.error(f"Error in map_line_to_param: {e}")
        return None, 0.0


VALUE_UNIT_RE = re.compile(
    r"([±]?\d+(?:\.\d+)?)\s*(mm|cm|m|µm|um|micron|bar|psi|°C|C|F)?",
    flags=re.IGNORECASE,
)


def extract_value_unit(text_line: str):
    m = VALUE_UNIT_RE.search(text_line)
    if not m:
        return None, None
    val = m.group(1)
    unit = m.group(2)
    if unit:
        unit = unit.replace("micron", "µm").lower()
    return val, unit


def normalize_numeric(value: str, unit: str, target="mm"):
    try:
        v = float(value.replace("±", ""))
    except Exception:
        return None
    if not unit:
        return str(v)
    u = unit.lower()
    if target == "mm":
        if u in ("mm",):
            return str(v)
        if u == "cm":
            return str(v * 10)
        if u == "m":
            return str(v * 1000)
        if u in ("µm", "um", "micron"):
            return str(v / 1000)
    if target == "um":
        if u in ("µm", "um", "micron"):
            return str(v)
        if u == "mm":
            return str(v * 1000)
    if target == "bar":
        if u in ("bar",):
            return str(v)
        if u == "psi":
            return str(v * 0.0689476)
    if target in ("c", "°c", "celsius"):
        if u in ("f",):
            return str((v - 32) * 5/9)
        return str(v)
    return str(v)


# helper to map extension to source string and priority rank
def source_type_and_priority(filepath: Path):
    suf = filepath.suffix.lower()
    if suf in (".docx",):
        return "DOCX", 1
    if suf in (".pdf",):
        return "PDF", 2
    if suf in (".png", ".jpg", ".jpeg", ".tiff"):
        return "Image", 3
    return "Other", 4


def process_all_and_build_master_from_s3(run_id: str, priority=("DOCX", "PDF", "Image")):
    """Process files directly from S3 without downloading to local storage."""
    session = SessionLocal()
    try:
        parsed_by_source = {}
        extraction_id_by_file = {}
        
        bucket = os.getenv("S3_BUCKET")
        if not bucket:
            raise RuntimeError("S3_BUCKET environment variable not set.")
        
        # Get all files from S3 for this run
        prefix = f"uploads/{run_id}/"
        s3_files = list(list_objects(bucket, prefix))
        
        if not s3_files:
            raise RuntimeError(f"No files found in S3 for run_id: {run_id}")
        
        logger.info(f"Processing {len(s3_files)} files from S3 for run_id: {run_id}")
        
        # --- Extract from S3 files ---
        for s3_key in s3_files:
            if s3_key.endswith("/"):  # Skip folders
                continue
                
            filename = s3_key.split("/")[-1]  # Get just the filename
            stype, sprio = source_type_and_priority(Path(filename))
            
            try:
                # Download file as stream from S3
                file_stream = download_file_stream(bucket, s3_key)
                raw_text = extract_text_for_s3_stream(file_stream, filename)
                logger.info(f"Extracted {len(raw_text)} chars from {filename} (type={stype}) from S3")
                
                # Debug: Log first few lines of extracted text
                if raw_text:
                    lines = raw_text.splitlines()[:5]
                    logger.info(f"First 5 lines from {filename}: {lines}")
                else:
                    logger.warning(f"No text extracted from {filename}")
                
                # Save raw extraction (DB)
                re_obj = RawExtraction(source=filename, raw_text=raw_text, meta_info={"type": stype, "s3_key": s3_key})
                session.add(re_obj)
                session.commit()
                extraction_id_by_file[filename] = re_obj.id
                logger.info(f"Saved raw extraction to DB with ID: {re_obj.id}")
                
                # Parse lines
                parsed = {}
                lines_processed = 0
                specs_found = 0
                
                for line in raw_text.splitlines():
                    if not line.strip():
                        continue
                    lines_processed += 1
                    
                    param, score = map_line_to_param(line)
                    if score < 0.55:
                        continue
                    
                    logger.info(f"Found potential param '{param}' with score {score:.3f} in line: {line.strip()}")
                    
                    val, unit = extract_value_unit(line)
                    if not val:
                        tokens = line.strip().split()
                        if len(tokens) >= 2:
                            candidate = tokens[-1]
                            val, unit = candidate, None
                    
                    if val:
                        specs_found += 1
                        # choose normalization target heuristically
                        target = None
                        if any(k in param for k in ("diameter", "hole", "cap", "thickness", "length", "width", "size")):
                            target = "mm"
                        elif "surface_finish" in param or "finish" in param:
                            target = "um"
                        elif "pressure" in param:
                            target = "bar"
                        elif "temperature" in param:
                            target = "C"

                        norm_val = normalize_numeric(val, unit, target) if target else val
                        parsed.setdefault(param, []).append({
                            "raw": line.strip(),
                            "value": norm_val,
                            "unit": unit or "",
                            "source": stype,
                            "priority": sprio,
                            "filename": filename,
                            "extraction_id": re_obj.id,
                            "s3_key": s3_key
                        })
                        
                        logger.info(f"Added spec: {param} = {norm_val} {unit or ''} (source: {stype})")
                
                logger.info(f"File {filename}: processed {lines_processed} lines, found {specs_found} specs")
                parsed_by_source[filename] = parsed
                
            except Exception as e:
                logger.error(f"Failed to process file {filename} from S3: {e}")
                continue
        
        return _build_master_from_parsed_data(session, parsed_by_source, extraction_id_by_file)
    finally:
        session.close()

def process_all_and_build_master(priority=("DOCX", "PDF", "Image")):
    """Legacy function for processing local files - kept for backward compatibility."""
    session = SessionLocal()
    parsed_by_source = {}
    extraction_id_by_file = {}

    # --- Extract from files ---
    for filepath in sorted(UPLOAD_DIR.iterdir()):
        if filepath.is_dir():
            continue
        stype, sprio = source_type_and_priority(filepath)
        raw_text = extract_text_for_file(filepath)
        logger.info(f"Extracted {len(raw_text)} chars from {filepath.name} (type={stype})")

        # save raw extraction (DB)
        landing_payload = {"source": filepath.name, "type": stype, "raw_text": raw_text}
        landing_path = LANDING_DIR / f"{filepath.name}.json"
        landing_path.write_text(json.dumps(landing_payload, ensure_ascii=False), encoding="utf-8")

        re_obj = RawExtraction(source=filepath.name, raw_text=raw_text, meta_info={"type": stype})
        session.add(re_obj)
        session.commit()
        extraction_id_by_file[filepath.name] = re_obj.id

        # parse lines
        parsed = {}
        for line in raw_text.splitlines():
            if not line.strip():
                continue
            param, score = map_line_to_param(line)
            if score < 0.55:
                continue
            val, unit = extract_value_unit(line)
            if not val:
                tokens = line.strip().split()
                if len(tokens) >= 2:
                    candidate = tokens[-1]
                    val, unit = candidate, None
            if val:
                # choose normalization target heuristically
                target = None
                if any(k in param for k in ("diameter", "hole", "cap", "thickness", "length", "width", "size")):
                    target = "mm"
                elif "surface_finish" in param or "finish" in param:
                    target = "um"
                elif "pressure" in param:
                    target = "bar"
                elif "temperature" in param:
                    target = "C"

                norm_val = normalize_numeric(val, unit, target) if target else val
                parsed.setdefault(param, []).append({
                    "raw": line.strip(),
                    "value": norm_val,
                    "unit": unit or "",
                    "source": stype,
                    "priority": sprio,
                    "filename": filepath.name,
                    "extraction_id": re_obj.id
                })

        parsed_by_source[filepath.name] = parsed
    
    return _build_master_from_parsed_data(session, parsed_by_source, extraction_id_by_file)

def _build_master_from_parsed_data(session, parsed_by_source, extraction_id_by_file):
    """Build master specifications from parsed data."""
    # --- Build in-memory master variants by param (list of variants) ---
    master_variants: Dict[str, List[Dict[str, Any]]] = {}
    for param in CANONICAL.keys():
        variants = []
        # collect from parsed_by_source: each source may have multiple lines for same param
        for filename, parsed_map in parsed_by_source.items():
            if not parsed_map:
                continue
            values_list = parsed_map.get(param, [])
            for v in values_list:
                variants.append(v)

        master_variants[param] = variants

    # --- Persist all variants into MasterSpec table (one row per variant) ---
    # We'll insert a row if exact (param, source, raw) doesn't already exist; else update.
    total_specs_to_save = 0
    specs_saved = 0
    
    for param, variants in master_variants.items():
        total_specs_to_save += len(variants)
        for variant in variants:
            # check if a row exists with same param + source + raw (very likely unique)
            existing = session.query(MasterSpec).filter(
                MasterSpec.param == param,
                MasterSpec.source == variant.get("source"),
                MasterSpec.raw == variant.get("raw")
            ).first()
            meta = {"filename": variant.get("filename"), "extraction_id": variant.get("extraction_id")}
            if variant.get("s3_key"):
                meta["s3_key"] = variant.get("s3_key")
            
            if existing:
                existing.value = variant.get("value")
                existing.unit = variant.get("unit")
                existing.priority = int(variant.get("priority", 10))
                existing.meta = meta
                logger.info(f"Updated existing spec: {param} = {variant.get('value')}")
                specs_saved += 1
            else:
                new = MasterSpec(
                    param=param,
                    value=variant.get("value"),
                    unit=variant.get("unit"),
                    raw=variant.get("raw"),
                    source=variant.get("source"),
                    priority=int(variant.get("priority", 10)),
                    meta=meta
                )
                session.add(new)
                logger.info(f"Added new spec to session: {param} = {variant.get('value')}")
                specs_saved += 1
    
    logger.info(f"Saving {specs_saved} specs to database (out of {total_specs_to_save} variants)")
    session.commit()
    logger.info("Successfully committed specs to database")

    # --- Build merged master: for each param, choose a 'chosen' variant by priority & recency ---
    merged_master = {}
    for param, variants in master_variants.items():
        # load persisted variants from DB for param to get added_at timestamps
        rows = session.query(MasterSpec).filter(MasterSpec.param == param).all()
        # Convert DB rows to variant records for API output
        api_variants = []
        for r in rows:
            api_variants.append({
                "id": r.id,
                "value": r.value,
                "unit": r.unit,
                "raw": r.raw,
                "source": r.source,
                "priority": r.priority,
                "meta": r.meta,
                "added_at": r.added_at.isoformat() if r.added_at else None
            })

        # choose one: lowest priority (i.e., priority int smaller) wins; tie-break by latest added_at
        chosen = None
        if api_variants:
            api_variants_sorted = sorted(api_variants, key=lambda x: (x.get("priority", 99), x.get("added_at") or ""))
            chosen = api_variants_sorted[0]
        merged_master[param] = {
            "chosen": chosen,
            "variants": api_variants
        }

    # --- Save CSV snapshot directly to S3 instead of local storage ---
    df_rows = []
    for p, info in merged_master.items():
        chosen = info.get("chosen") or {}
        df_rows.append({"param": p, "value": chosen.get("value") or "", "unit": chosen.get("unit") or "", "source": chosen.get("source") or ""})
    df_out = pd.DataFrame(df_rows)
    
    # Save to temporary file and upload to S3
    with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False) as temp_file:
        df_out.to_csv(temp_file.name, index=False)
        temp_file_path = temp_file.name
    
    # Upload to S3
    bucket = os.getenv("S3_BUCKET")
    if bucket:
        s3_key = f"outputs/master_specs.csv"
        upload_file(bucket, temp_file_path, s3_key)
        logger.info(f"Master specs CSV uploaded to s3://{bucket}/{s3_key}")
    
    # Clean up temp file
    os.unlink(temp_file_path)

    return parsed_by_source, merged_master


# ========== defect mapping (unchanged semantics but now expects merged_master format) ==========
def classify_defect_with_master(defect: dict, merged_master: dict):
    dtype = (defect.get("defect_type") or "").lower()
    
    # Find all rules for this defect type
    rules = [r for r in DEFECT_RULES if (r.get("defect_type") or "").lower() == dtype]
    if not rules:
        return "Unknown1"

    # We'll assume one rule per defect type for simplicity
    rule = rules[0]

    # Special cases
    special = rule.get("special", "").lower()
    if special == "always_fail":
        return rule.get("fail", "Not Repairable")
    if special == "coating":
        coating_info = merged_master.get("coating_required", {}).get("chosen")
        if coating_info and str(coating_info.get("value")).lower() in ("yes", "true", "1"):
            return rule.get("fail", "Not Repairable")
        return rule.get("ok", "Repairable")

    spec_name = rule.get("spec_name")
    field = rule.get("field")
    op = rule.get("op")

    spec_info = merged_master.get(spec_name, {}).get("chosen") if spec_name else None
    if spec_name and (not spec_info or not spec_info.get("value")):
        return "Unknown2"

    spec_val = spec_info["value"] if spec_info else None
    field_val = defect.get(field) if field else None

    # Numeric comparison
    try:
        if spec_val is not None and field_val is not None and op in ("<=", "<", ">=", ">", "=="):
            spec_num = float(spec_val)
            val_num = float(field_val)
            if eval(f"{val_num} {op} {spec_num}"):
                return rule.get("ok", "Repairable")
            else:
                return rule.get("fail", "Not Repairable")
    except Exception:
        # fallback to string comparison
        if op == "==" and field_val is not None and spec_val is not None:
            return rule.get("ok") if str(field_val).lower() == str(spec_val).lower() else rule.get("fail")

    return "Unknown3"


def run_defect_mapping(defect_file_path: Path, merged_master: Dict[str, Any]) -> pd.DataFrame:
    """
    Run defect mapping and return results as a DataFrame.
    Saves defect_results.csv to S3 instead of local storage.
    """
    # Load defects CSV or JSON
    if defect_file_path.suffix.lower() == ".csv":
        df = pd.read_csv(defect_file_path)
    else:
        records = json.loads(defect_file_path.read_text())
        df = pd.DataFrame(records)

    # Apply defect classification
    df["decision"] = df.apply(lambda r: classify_defect_with_master(r.to_dict(), merged_master), axis=1)

    # Save results to S3 instead of local storage
    bucket = os.getenv("S3_BUCKET")
    if bucket:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False) as temp_file:
            df.to_csv(temp_file.name, index=False)
            temp_file_path = temp_file.name
        
        s3_key = f"outputs/defect_results.csv"
        upload_file(bucket, temp_file_path, s3_key)
        logger.info(f"Defect results CSV uploaded to s3://{bucket}/{s3_key}")
        
        # Clean up temp file
        os.unlink(temp_file_path)
    else:
        # Fallback to local storage if S3 not configured
        out_path = Path("data/outputs/defect_results.csv")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        df.to_csv(out_path, index=False)

    return df


def pull_uploads_from_s3(run_id: str = "latest"):
    """
    Downloads S3 prefix s3://{S3_BUCKET}/uploads/{run_id}/ into UPLOAD_DIR (clears UPLOAD_DIR first).
    """
    bucket = os.getenv("S3_BUCKET")
    if not bucket:
        raise RuntimeError("S3_BUCKET environment variable not set. Please configure S3_BUCKET in your environment.")

    try:
        logger.info(f"Pulling uploads from S3: s3://{bucket}/uploads/{run_id}/")
        
        # Optionally clear existing uploads to avoid duplicates
        if UPLOAD_DIR.exists():
            for f in UPLOAD_DIR.iterdir():
                if f.is_file():
                    f.unlink()
                else:
                    shutil.rmtree(f)
        UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

        prefix = f"uploads/{run_id}/"
        download_prefix(bucket, prefix, str(UPLOAD_DIR))
        
        # Check if any files were downloaded
        files_in_upload_dir = list(UPLOAD_DIR.iterdir())
        if not files_in_upload_dir:
            logger.warning(f"No files found in s3://{bucket}/uploads/{run_id}/")
        else:
            logger.info(f"Successfully pulled {len(files_in_upload_dir)} files from S3")
            
    except Exception as e:
        logger.error(f"Failed to pull uploads from S3: {e}")
        raise RuntimeError(f"S3 download failed: {e}")


def push_outputs_to_s3(run_id: str = None):
    """
    Upload outputs folder to s3://{S3_BUCKET}/outputs/{run_id or timestamp}/
    """
    bucket = os.getenv("S3_BUCKET")
    if not bucket:
        raise RuntimeError("S3_BUCKET environment variable not set. Please configure S3_BUCKET in your environment.")
    
    try:
        if run_id is None:
            run_id = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
        
        logger.info(f"Pushing outputs to S3: s3://{bucket}/outputs/{run_id}/")
        
        # Check if OUTPUT_DIR has any files to upload
        if not OUTPUT_DIR.exists() or not any(OUTPUT_DIR.iterdir()):
            logger.warning(f"No output files found in {OUTPUT_DIR} to upload to S3")
            return
        
        prefix = f"outputs/{run_id}"
        upload_folder(bucket, str(OUTPUT_DIR), prefix)
        
        logger.info(f"Successfully pushed outputs to s3://{bucket}/outputs/{run_id}/")
        
    except Exception as e:
        logger.error(f"Failed to push outputs to S3: {e}")
        raise RuntimeError(f"S3 upload failed: {e}")